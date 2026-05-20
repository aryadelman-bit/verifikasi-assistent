from docx import Document
from io import BytesIO
from utils.generator import generate_return_text

def create_docx():
    doc = Document()
    doc.add_heading('Hasil Verifikasi Dokumen Standar Kegiatan Usaha', 0)
    
    text = generate_return_text()
    
    for line in text.split('\n'):
        line_str = line.strip()
        if not line_str:
            continue
            
        # Jika depannya adalah strip, jadikan bullet list
        if line_str.startswith('- '):
            doc.add_paragraph(line_str[2:], style='List Bullet')
        # Jika merupakan pembuka/penutup surat, biarkan normal
        elif "Yth." in line_str or "Berkenaan dengan" in line_str or "Demikian disampaikan" in line_str or line_str.startswith("*"):
            doc.add_paragraph(line_str)
        # Sisa text lainnya dianggap sebagai Judul Section/Lampiran (di-Bold otomatis)
        else:
            p = doc.add_paragraph()
            p.add_run(line_str).bold = True
            
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream